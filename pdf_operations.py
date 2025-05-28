# Major-Project/pdf_operations.py


import os
import re
import io
import zipfile
import fitz 
from fitz.utils import getColor
import os 
from pathlib import Path 
import warnings
import logging
import subprocess
from datetime import datetime
from pathlib import Path
from pypdf import PdfReader, PdfWriter
from pypdf.errors import PdfReadError, FileNotDecryptedError
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches

try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    logging.warning("python-docx library not found. PDF-to-Word functionality will be disabled.")
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

# --- Configuration ---
OUTPUT_DIR = Path("output")
POPPLER_PATH = os.environ.get('POPPLER_PATH', None)
warnings.filterwarnings("ignore", category=UserWarning, module='pypdf')

# --- Helper Functions ---
def ensure_output_dir():
    """Creates the output directory if it doesn't exist."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    # logger.info(f"Ensured output directory exists: {OUTPUT_DIR.resolve()}") # Optional logging

def get_output_filename(base_name, suffix, extension):
    """Generates a unique output filename in the OUTPUT_DIR."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S%f") # Added microseconds
    safe_base = "".join(c if c.isalnum() or c in ('_', '-') else '_' for c in str(base_name))
    max_base_len = 100
    safe_base = safe_base[:max_base_len]
    filename = f"{safe_base}_{suffix}_{timestamp}{extension}"
    return OUTPUT_DIR / filename

# --- Core PDF Operations ---


def merge_pdfs(pdf_files, output_filename_base="merged"):
    """Merges multiple PDF file streams into one."""
    ensure_output_dir()
    merger = PdfWriter()
    processed_count = 0
    try:
        for pdf_stream in pdf_files:
            filename_for_log = getattr(pdf_stream, 'filename', 'N/A')
            try:
                reader = PdfReader(pdf_stream)
                if reader.is_encrypted:
                    try:
                        if reader.decrypt('') == 0: # Try empty password
                             logger.warning(f"Skipping encrypted file (password needed): {filename_for_log}")
                             continue # Skip this file
                    except FileNotDecryptedError:
                        logger.warning(f"Skipping encrypted file (password needed): {filename_for_log}")
                        continue # Skip this file
                    except Exception as decrypt_err:
                         logger.warning(f"Skipping file {filename_for_log} due to decryption check error: {decrypt_err}")
                         continue

                merger.append(reader)
                processed_count += 1
                # logger.info(f"Appended '{filename_for_log}' to merge.") # Verbose logging
            except PdfReadError as read_err:
                 logger.error(f"Error reading PDF stream '{filename_for_log}': {read_err}. Skipping.")
                 continue # Skip invalid PDF
            except Exception as e:
                logger.error(f"Unexpected error processing stream '{filename_for_log}' for merge: {e}. Skipping.", exc_info=True)
                continue # Skip on other errors

        if processed_count < 1:
            return None, "No valid PDF files could be processed for merging."

        output_path = get_output_filename(output_filename_base, "merged", ".pdf")
        with open(output_path, "wb") as f_out:
            merger.write(f_out)
        merger.close()
        logger.info(f"Successfully merged {processed_count} files into {output_path}")
        return output_path, None
    except Exception as e:
        logger.error(f"Error during final merge write operation: {e}", exc_info=True)
        if merger: merger.close()
        return None, f"Error finalizing merged PDF: {e}"



def parse_page_ranges(ranges_str, total_pages):
    """Parses a range string (e.g., '1-3, 5, 8-') into a list of tuples.
       Each tuple contains: (range_string_part, list_of_0_based_indices).
       Returns None, error_message if parsing fails.
    """
    if not ranges_str:
        return None, "Page range string cannot be empty."

    parsed_ranges = []
    parts = ranges_str.split(',')

    for part in parts:
        part = part.strip()
        if not part: continue

        indices = set()
        try:
            if '-' in part:
                start_str, end_str = part.split('-', 1)
                start = int(start_str) if start_str else 1
                end = int(end_str) if end_str else total_pages
                if not (1 <= start <= end <= total_pages):
                    raise ValueError(f"Range '{part}' is invalid for page count {total_pages}.")
                indices.update(range(start - 1, end))
            else:
                page_num = int(part)
                if not (1 <= page_num <= total_pages):
                    raise ValueError(f"Page number '{part}' is out of bounds (1-{total_pages}).")
                indices.add(page_num - 1)

            if indices:
                # Sanitize part string for use in filenames
                safe_part_str = re.sub(r'[^\w\-]+', '_', part)
                parsed_ranges.append((safe_part_str, sorted(list(indices))))

        except ValueError as ve:
            logger.error(f"Invalid page range format: {ve}")
            return None, f"Invalid page range format: {ve}"

    if not parsed_ranges:
        return None, "No valid pages or ranges specified."

    return parsed_ranges, None

# (Paste the split_pdf_to_multiple_files function from previous answer here)
def split_pdf_to_multiple_files(pdf_file_stream, ranges_str, output_filename_base="split"):
    """Splits a PDF based on page ranges into MULTIPLE output PDF files.
       Returns a list of output file paths, or None and an error message.
    """
    ensure_output_dir()
    output_paths = []
    reader = None
    filename_for_log = getattr(pdf_file_stream, 'filename', 'N/A')

    try:
        # Ensure stream is at the beginning before reading
        pdf_file_stream.seek(0)
        reader = PdfReader(pdf_file_stream)
        pdf_file_stream.seek(0) # Reset stream just in case reader consumes it

        if reader.is_encrypted:
            try:
                # Try decrypting with empty password. If it requires a password, fail.
                if reader.decrypt('') == 0:
                    err_msg = f"Error: Input PDF '{filename_for_log}' is password protected."
                    logger.error(err_msg)
                    return None, err_msg
            except FileNotDecryptedError:
                 err_msg = f"Error: Input PDF '{filename_for_log}' is password protected."
                 logger.error(err_msg)
                 return None, err_msg

        total_pages = len(reader.pages)
        logger.info(f"Processing multi-split for '{filename_for_log}' ({total_pages} pages) with ranges '{ranges_str}'.")

        # Use the parser helper function
        parsed_ranges, error_msg = parse_page_ranges(ranges_str, total_pages)
        if error_msg:
            return None, error_msg
        
        if parsed_ranges and len(parsed_ranges) > 10:
            limit_err = "Error: Splitting is limited to a maximum of 10 output files per request."
            logger.warning(f"{limit_err} Requested ranges would create {len(parsed_ranges)} files.")
            return None, limit_err
        
        if not parsed_ranges:
            logger.warning("No valid ranges resulted in pages to extract.")
            return [], None # Return empty list if no ranges valid

        # Loop through each parsed range part
        for range_label, indices in parsed_ranges:
            if not indices: continue # Skip empty index lists

            writer = PdfWriter()
            logger.info(f"Creating split file for range '{range_label}' with pages (0-based): {indices}")
            for index in indices:
                try:
                    # Add page *from the original reader*
                    writer.add_page(reader.pages[index])
                except IndexError:
                     logger.error(f"Page index {index} out of bounds for PDF {filename_for_log}. Skipping page.")
                     continue # Skip invalid index

            if not writer.pages:
                logger.warning(f"No valid pages added for range '{range_label}'. Skipping file creation.")
                continue # Don't create an empty PDF

            split_suffix = f"split_{range_label}"
            output_path = get_output_filename(output_filename_base, split_suffix, ".pdf")

            try:
                with open(output_path, "wb") as f_out:
                    writer.write(f_out)
                writer.close()
                output_paths.append(output_path)
                logger.info(f"Created split file: {output_path}")
            except Exception as write_err:
                logger.error(f"Failed to write split file for range {range_label}: {write_err}", exc_info=True)
                if writer: writer.close()
                # Cleanup already created files before returning error
                for p in output_paths:
                     try: os.remove(p)
                     except OSError: pass
                return None, f"Error writing split file for range {range_label}: {write_err}"

        logger.info(f"Successfully created {len(output_paths)} split PDF file(s).")
        return output_paths, None

    except (PdfReadError, ValueError, Exception) as e:
        logger.error(f"Error during multi-split PDF process for {filename_for_log}: {e}", exc_info=True)
        # Cleanup any files created before the error
        for p in output_paths:
             try: os.remove(p)
             except OSError: pass
        return None, f"Error splitting PDF: {e}"

# --- END SPLIT PDF ---


# --- ROTATE PDF ---
# (Paste the working rotate_pdf function definition from previous answer here)
def rotate_pdf(pdf_file_stream, rotation_angle, output_filename_base="rotated"):
    """Rotates all pages in a PDF stream by a specified angle (90, 180, 270)."""
    ensure_output_dir()
    if rotation_angle not in [90, 180, 270]:
        logger.error(f"Invalid rotation angle specified: {rotation_angle}")
        return None, "Error: Rotation angle must be 90, 180, or 270."

    writer = PdfWriter()
    filename_for_log = getattr(pdf_file_stream, 'filename', 'N/A')
    try:
        pdf_file_stream.seek(0) # Ensure stream is at start
        reader = PdfReader(pdf_file_stream)
        pdf_file_stream.seek(0) # Reset

        if reader.is_encrypted:
            try:
                if reader.decrypt('') == 0:
                    logger.error(f"Cannot rotate password-protected PDF: {filename_for_log}")
                    return None, f"Error: Input PDF '{filename_for_log}' is password protected."
            except FileNotDecryptedError:
                logger.error(f"Cannot rotate password-protected PDF: {filename_for_log}")
                return None, f"Error: Input PDF '{filename_for_log}' is password protected."

        logger.info(f"Rotating {len(reader.pages)} pages in {filename_for_log} by {rotation_angle} degrees.")
        for page in reader.pages:
            page.rotate(rotation_angle)
            writer.add_page(page)

        output_path = get_output_filename(output_filename_base, f"rotated_{rotation_angle}", ".pdf")
        with open(output_path, "wb") as f_out:
            writer.write(f_out)
        writer.close()
        logger.info(f"Rotated PDF saved to: {output_path}")
        return output_path, None

    except Exception as e:
        logger.error(f"Error rotating PDF {filename_for_log}: {e}", exc_info=True)
        if writer: writer.close()
        return None, f"Error rotating PDF: {e}"
# --- END ROTATE PDF ---


# --- PROTECT PDF ---
# (Paste the working add_password function definition from previous answer here)
def add_password(pdf_file_stream, password, output_filename_base="protected"):
    """Adds a user password to encrypt the PDF stream."""
    ensure_output_dir()
    if not password:
        logger.error("Password cannot be empty for protection.")
        return None, "Error: Password cannot be empty."

    writer = PdfWriter()
    filename_for_log = getattr(pdf_file_stream, 'filename', 'N/A')
    try:
        pdf_file_stream.seek(0) # Ensure stream is at start
        reader = PdfReader(pdf_file_stream)
        pdf_file_stream.seek(0) # Reset

        if reader.is_encrypted:
            logger.warning(f"Input PDF {filename_for_log} is already password protected.")
            return None, "Error: Input PDF is already password protected."

        for page in reader.pages:
            writer.add_page(page)

        logger.info(f"Encrypting PDF {filename_for_log} with AES-256.")
        writer.encrypt(user_password=password, algorithm="AES-256") # Requires 'cryptography' package

        output_path = get_output_filename(output_filename_base, "protected", ".pdf")
        with open(output_path, "wb") as f_out:
            writer.write(f_out)
        writer.close()
        logger.info(f"Successfully protected PDF saved to: {output_path}")
        return output_path, None
    except ImportError as imp_err:
         # Specifically catch if cryptography is missing
         if 'cryptography' in str(imp_err):
             logger.error("Cryptography package needed for AES encryption is not installed. Install it via pip.", exc_info=True)
             return None, "Error: Required 'cryptography' library not found. Cannot protect with AES."
         else:
              logger.error(f"Import error during password addition for {filename_for_log}: {imp_err}", exc_info=True)
              if writer: writer.close()
              return None, f"Error adding password (import issue): {imp_err}"
    except Exception as e:
        logger.error(f"Error adding password to {filename_for_log}: {e}", exc_info=True)
        if writer: writer.close()
        return None, f"Error adding password: {e}"
# --- END PROTECT PDF ---


# --- UNLOCK PDF ---
# (Paste a working remove_password function definition here, ensure logging)
def remove_password(pdf_file_stream, password, output_filename_base="unlocked"):
    """Removes the password from a PDF stream, given the correct password."""
    ensure_output_dir()
    if not password:
        return None, "Error: Password needed to unlock."

    writer = PdfWriter()
    filename_for_log = getattr(pdf_file_stream, 'filename', 'N/A')
    try:
        pdf_file_stream.seek(0) # Ensure stream is at start
        reader = PdfReader(pdf_file_stream)
        pdf_file_stream.seek(0) # Reset

        if not reader.is_encrypted:
            logger.warning(f"PDF {filename_for_log} is not password protected.")
            return None, "Error: PDF is not password protected."

        logger.info(f"Attempting to decrypt {filename_for_log}...")
        if reader.decrypt(password) == 0: # 0 means decryption failed
             logger.error(f"Incorrect password provided for {filename_for_log}")
             return None, "Error: Incorrect password provided."
        # Decryption successful if it returns 1 or 2
        logger.info(f"Decryption successful for {filename_for_log}.")

        # Clone the decrypted content
        writer.clone_document_from_reader(reader)

        output_path = get_output_filename(output_filename_base, "unlocked", ".pdf")
        with open(output_path, "wb") as f_out:
            writer.write(f_out)
        writer.close()
        logger.info(f"Unlocked PDF saved to: {output_path}")
        return output_path, None

    except FileNotDecryptedError: # Catch specifically if decrypt fails strongly
         logger.error(f"Incorrect password provided (FileNotDecryptedError) for {filename_for_log}")
         return None, "Error: Incorrect password provided."
    except Exception as e:
        logger.error(f"Error removing password for {filename_for_log}: {e}", exc_info=True)
        if writer: writer.close()
        return None, f"Error removing password: {e}"
# --- END UNLOCK PDF ---


# --- PDF TO IMAGES ---
# (Paste the pdf_to_images function from previous answer here, ensure logging)
def pdf_to_images(pdf_file, fmt='jpeg', dpi=200, output_filename_base="page"):
    """Converts each page of a PDF (path or stream) to image files."""
    ensure_output_dir()
    output_paths = []
    fmt = fmt.lower()
    if fmt not in ['jpeg', 'png']:
        return [], "Error: Unsupported image format. Use 'jpeg' or 'png'."

    temp_pdf_path_obj = None # Use Path object for consistency
    input_path_str = None # Path as string for pdf2image
    filename_for_log = getattr(pdf_file, 'filename', 'N/A')

    try:
        # Handle stream input by saving temporarily
        if isinstance(pdf_file, (io.BytesIO, io.BufferedReader)):
            temp_pdf_path_obj = OUTPUT_DIR / f"temp_img_conv_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.pdf"
            logger.info(f"Input is a stream for pdf_to_images, saving temporarily to {temp_pdf_path_obj}")
            with open(temp_pdf_path_obj, 'wb') as f:
                pdf_file.seek(0)
                f.write(pdf_file.read())
                pdf_file.seek(0)
            input_path_str = str(temp_pdf_path_obj)
            filename_for_log = Path(filename_for_log).stem # Use stem from original name if possible
        elif isinstance(pdf_file, (str, Path)):
             input_path_str = str(pdf_file)
             filename_for_log = Path(input_path_str).stem
        else:
             raise TypeError("Unsupported input type for pdf_file. Must be path string or stream.")

        logger.info(f"Attempting to convert PDF '{Path(input_path_str).name}' to {fmt} images (DPI: {dpi}).")
        logger.info(f"Using Poppler path: {POPPLER_PATH or 'System PATH'}")

        # Check for encryption *before* passing to pdf2image
        try:
            reader_check = PdfReader(input_path_str)
            if reader_check.is_encrypted:
                if reader_check.decrypt('') == 0:
                     err_msg = f"Error: Input PDF '{Path(input_path_str).name}' is password protected."
                     logger.error(err_msg)
                     # Cleanup temp file before returning
                     if temp_pdf_path_obj and temp_pdf_path_obj.exists(): os.remove(temp_pdf_path_obj)
                     return [], err_msg
        except Exception as pdf_err:
            logger.error(f"Error checking PDF encryption for '{Path(input_path_str).name}': {pdf_err}")
            if temp_pdf_path_obj and temp_pdf_path_obj.exists(): os.remove(temp_pdf_path_obj)
            return [], f"Error reading input PDF: {pdf_err}"

        # Convert using pdf2image
        images = convert_from_path(
            input_path_str,
            dpi=dpi,
            fmt=fmt,
            poppler_path=POPPLER_PATH,
            thread_count=2
        )

        if not images:
            logger.error("pdf2image returned no images. Check Poppler installation and PATH.")
            if temp_pdf_path_obj and temp_pdf_path_obj.exists(): os.remove(temp_pdf_path_obj)
            return [], "Error converting PDF to images. Check Poppler installation/PATH."

        logger.info(f"Successfully generated {len(images)} image objects from PDF.")
        ext = ".jpg" if fmt == 'jpeg' else ".png"

        for i, image in enumerate(images):
            output_path = get_output_filename(filename_for_log, f"page_{i+1}", ext)
            try:
                 # Handle potential transparency for PNGs before saving JPEG
                 if fmt == 'jpeg' and image.mode in ('RGBA', 'LA', 'P'):
                      logger.debug(f"Converting image {i+1} to RGB before saving as JPEG.")
                      # Create a white background image
                      bg = Image.new("RGB", image.size, (255, 255, 255))
                      # Paste the image onto the background using its alpha channel or P mode palette
                      bg.paste(image, (0,0), image if image.mode == 'RGBA' or image.mode == 'LA' else None)
                      image_to_save = bg
                 else:
                      image_to_save = image

                 image_to_save.save(output_path, fmt.upper())
                 output_paths.append(output_path)
                 # logger.debug(f"Saved image: {output_path}") # Verbose
            except Exception as save_err:
                logger.error(f"Failed to save image {i+1} to {output_path}: {save_err}", exc_info=True)
                # Cleanup created images and temp file
                for p in output_paths: os.remove(p)
                if temp_pdf_path_obj and temp_pdf_path_obj.exists(): os.remove(temp_pdf_path_obj)
                return [], f"Error saving generated image: {save_err}"

        logger.info(f"Successfully saved {len(output_paths)} images.")
        return output_paths, None

    except Exception as e:
        logger.error(f"Error converting PDF '{filename_for_log}' to images: {e}", exc_info=True)
        if "pdfinfo" in str(e) or "pdftoppm" in str(e) or "Poppler" in str(e):
             err_msg = f"Error during conversion, likely Poppler related. Is Poppler installed and in PATH? Details: {e}"
        else:
             err_msg = f"Unexpected error converting PDF to images: {e}"
        # Cleanup temp file if it exists
        if temp_pdf_path_obj and temp_pdf_path_obj.exists(): os.remove(temp_pdf_path_obj)
        return [], err_msg
    finally:
        # Final cleanup check for temp file
        if temp_pdf_path_obj and temp_pdf_path_obj.exists():
            try:
                os.remove(temp_pdf_path_obj)
                logger.info(f"Removed temporary file: {temp_pdf_path_obj}")
            except OSError as e_rem:
                logger.warning(f"Could not remove temporary file {temp_pdf_path_obj}: {e_rem}")
# --- END PDF TO IMAGES ---


# --- IMAGE TO PDF ---
# (Paste a working images_to_pdf function definition here, ensure logging)
def images_to_pdf(image_files, output_filename_base="from_images"):
    """Converts multiple image file streams into a single PDF."""
    ensure_output_dir()
    pil_images = []
    processed_files_info = [] # Store filenames for logging

    try:
        for img_stream in image_files:
            filename = getattr(img_stream, 'filename', 'N/A')
            try:
                img_stream.seek(0) # Reset stream
                img = Image.open(img_stream)
                # Convert common non-RGB modes to RGB for broader PDF compatibility
                if img.mode == 'RGBA' or img.mode == 'P' or img.mode == 'LA':
                    logger.debug(f"Converting image '{filename}' from {img.mode} to RGB.")
                    # Create a white background and paste image with alpha mask if applicable
                    if img.mode == 'RGBA' or img.mode == 'LA':
                         alpha = img.split()[-1]
                         bg = Image.new("RGB", img.size, (255, 255, 255))
                         bg.paste(img, mask=alpha)
                         img_converted = bg
                    else: # Palette mode, P
                         img_converted = img.convert('RGB')
                    pil_images.append(img_converted)
                elif img.mode == 'RGB':
                    pil_images.append(img) # Already RGB
                else:
                     # Attempt conversion for other modes like L (grayscale), CMYK etc.
                     logger.debug(f"Attempting to convert image '{filename}' from {img.mode} to RGB.")
                     img_converted = img.convert('RGB')
                     pil_images.append(img_converted)

                processed_files_info.append(filename)
            except Exception as e:
                logger.warning(f"Skipping file {filename} due to error opening or converting image: {e}")
                continue # Skip this image

        if not pil_images:
            return None, "Error: No valid images found or processed."

        output_path = get_output_filename(output_filename_base, "converted", ".pdf")
        logger.info(f"Converting {len(pil_images)} images ({', '.join(processed_files_info)}) to PDF: {output_path}")

        # Save the first image, then append the rest
        pil_images[0].save(
            output_path,
            "PDF",
            resolution=100.0,
            save_all=True,
            append_images=pil_images[1:]
        )

        return output_path, None
    except Exception as e:
        logger.error(f"Error converting images to PDF: {e}", exc_info=True)
        return None, f"Error converting images to PDF: {e}"
    finally:
        # Close images (optional but good practice)
        for img in pil_images:
            try:
                img.close()
            except Exception:
                pass
# --- END IMAGE TO PDF ---


# --- OFFICE TO PDF ---
# (Paste the office_to_pdf function from previous answer here, ensure logging/path handling)
def office_to_pdf(office_file_path, output_filename_base="converted"):
    """Converts an Office document (Word, Excel, PPT) to PDF using LibreOffice."""
    ensure_output_dir()
    output_dir_abs = OUTPUT_DIR.resolve() # LibreOffice needs an absolute path
    input_file_path = Path(office_file_path).resolve() # Ensure input is absolute path too
    input_filename = input_file_path.name

    logger.info(f"Attempting to convert Office file '{input_filename}' to PDF using LibreOffice.")

    # --- Find soffice ---
    soffice_command = os.environ.get('SOFFICE_PATH') # Prioritize environment variable
    if soffice_command and Path(soffice_command).is_file(): # Check if it's a file
         logger.info(f"Using soffice path from SOFFICE_PATH env var: {soffice_command}")
    else:
        # Search in common locations within the container/system
        possible_paths = ["soffice", "libreoffice", "/usr/bin/soffice", "/usr/bin/libreoffice"]
        found = False
        for cmd_path in possible_paths:
            try:
                logger.debug(f"Checking for soffice at: {cmd_path}")
                result = subprocess.run([cmd_path, '--version'], check=True, capture_output=True, text=True, timeout=10)
                logger.info(f"Found working LibreOffice command: {cmd_path}. Version: {result.stdout.strip()}")
                soffice_command = cmd_path
                found = True
                break
            except (subprocess.CalledProcessError, FileNotFoundError, PermissionError, subprocess.TimeoutExpired) as check_err:
                logger.debug(f"Checking '{cmd_path}' failed: {check_err}")
                continue
        if not found:
             msg = "Error: LibreOffice 'soffice' command not found or not executable in expected paths. Install LibreOffice or set SOFFICE_PATH."
             logger.error(msg)
             return None, msg
    # --- End Find soffice ---

    try:
        cmd = [
            soffice_command,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(output_dir_abs),
            str(input_file_path)
        ]
        logger.info(f"Running LibreOffice command: {' '.join(cmd)}")
        result = subprocess.run(cmd, check=True, capture_output=True, text=True, timeout=300) # 5 min timeout
        logger.info(f"LibreOffice stdout: {result.stdout or '[No stdout]'}")
        if result.stderr:
            logger.warning(f"LibreOffice stderr: {result.stderr}")

        # --- Verify output ---
        original_stem = input_file_path.stem
        expected_pdf = output_dir_abs / f"{original_stem}.pdf"

        if expected_pdf.exists():
             logger.info(f"LibreOffice successfully created: {expected_pdf}")
             final_output_path = get_output_filename(output_filename_base or original_stem, "from_office", ".pdf")
             try:
                 expected_pdf.rename(final_output_path)
                 logger.info(f"Renamed output file to: {final_output_path}")
                 return final_output_path, None
             except OSError as rename_err:
                 logger.error(f"Failed to rename '{expected_pdf}' to '{final_output_path}': {rename_err}")
                 return expected_pdf, None # Return original path if rename fails

        else:
            err_msg = f"Error: Expected PDF '{expected_pdf}' not found after LibreOffice command."
            logger.error(err_msg)
            if result.stderr and "Error:" in result.stderr:
                 err_detail = f" Details: {result.stderr[:500]}"
                 return None, err_msg + err_detail
            else:
                 return None, err_msg + " Check LibreOffice compatibility or installation."

    except FileNotFoundError:
         msg = f"Error: LibreOffice command '{soffice_command}' not found during execution."
         logger.error(msg)
         return None, msg
    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice conversion failed (exit code {e.returncode}). Stderr: {e.stderr}", exc_info=True)
        err_msg = f"Error during LibreOffice conversion (Code {e.returncode})."
        if e.stderr: err_msg += f" Details: {e.stderr[:500]}"
        return None, err_msg
    except subprocess.TimeoutExpired:
        logger.error("Error: LibreOffice conversion timed out.")
        return None, "Error: Office to PDF conversion timed out (> 5 minutes)."
    except Exception as e:
        logger.error(f"Unexpected error during Office to PDF conversion: {e}", exc_info=True)
        return None, f"Unexpected error during Office to PDF conversion: {e}"
# --- END OFFICE TO PDF ---


def compress_pdf(pdf_path_or_stream, output_filename_base="compressed", **kwargs): # Add **kwargs
    """Compresses a PDF file using PyMuPDF optimizations."""
    ensure_output_dir()
    doc = None
    filename_for_log = "input_stream"
    temp_pdf_path_obj = None
    original_size = 0
    compressed_size = 0

    try:
        # --- Determine original size ---
        if isinstance(pdf_path_or_stream, (str, Path)):
            pdf_path = str(pdf_path_or_stream)
            filename_for_log = Path(pdf_path).name
            original_size = Path(pdf_path).stat().st_size
            doc = fitz.open(pdf_path)
        elif isinstance(pdf_path_or_stream, (io.BytesIO, io.BufferedReader)):
            filename_for_log = getattr(pdf_path_or_stream, 'filename', 'input_stream')
            
            # Read stream to get its size
            pdf_path_or_stream.seek(0)
            stream_content = pdf_path_or_stream.read()
            original_size = len(stream_content)
            pdf_path_or_stream.seek(0) # Reset stream

            # Save stream temporarily as fitz.open might need path for some operations or complex PDFs
            temp_dir = OUTPUT_DIR 
            temp_pdf_path_obj = temp_dir / f"temp_compress_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.pdf"
            logger.info(f"Input is a stream for compression, saving temporarily to {temp_pdf_path_obj}")
            with open(temp_pdf_path_obj, 'wb') as f:
                f.write(stream_content) # Write the read content
            doc = fitz.open(str(temp_pdf_path_obj))
            filename_for_log = Path(filename_for_log).stem 
        else:
            raise TypeError("Unsupported input type for compress_pdf. Must be path string or stream.")

        if doc.is_encrypted:
            logger.error(f"Cannot compress password-protected PDF: {filename_for_log}")
            if doc: doc.close()
            if temp_pdf_path_obj: cleanup_temp_file(temp_pdf_path_obj)
            # Return None for sizes as well
            return None, f"Error: Input PDF '{filename_for_log}' is password protected.", None, None 

        logger.info(f"Compressing PDF '{filename_for_log}' using PyMuPDF...")
        output_path = get_output_filename(output_filename_base or Path(filename_for_log).stem, "compressed", ".pdf")

        # PyMuPDF save options for compression:
        # Default options from previous state, will be augmented by kwargs
        save_params = {'garbage': 4, 'deflate': True}
        save_params.update(kwargs) # Apply any new options passed

        doc.save(str(output_path), **save_params) # Use the merged parameters
        doc.close() 

        compressed_size = output_path.stat().st_size
        reduction_percent = 0
        if original_size > 0 and compressed_size > 0: # ensure compressed_size is also positive
            reduction_percent = (1 - compressed_size / original_size) * 100
            logger.info(f"Compression complete: {output_path}. Original: {original_size}B, Compressed: {compressed_size}B, Reduction: {reduction_percent:.1f}%")
        elif original_size > 0 :
            logger.info(f"Compression complete: {output_path}. Original: {original_size}B, Compressed: {compressed_size}B. Could not calculate reduction percentage.")
        else:
            logger.info(f"Compression complete: {output_path}. Compressed size: {compressed_size} bytes. Original size unknown or zero.")
        
        # Return output_path, error_msg (None on success), original_size, compressed_size
        return output_path, None, original_size, compressed_size

    except Exception as e:
        logger.error(f"Error compressing PDF '{filename_for_log}': {e}", exc_info=True)
        if doc: doc.close()
        return None, f"Error compressing PDF: {e}", original_size, 0 # Return original_size and 0 for compressed
    finally:
        if temp_pdf_path_obj:
            cleanup_temp_file(temp_pdf_path_obj)
# --- END COMPRESS PDF ---


def pdf_to_word(pdf_path_or_stream, output_filename_base="converted"):
    """
    Converts PDF to a Word (.docx) file, extracting text and basic image layout.
    NOTE: Complex formatting (tables, columns, vector graphics, precise styling) WILL BE LOST.
    """
    if not DOCX_AVAILABLE:
        logger.error("Can't do PDF-to-Word: python-docx isn't installed.")
        return None, "Error: Required 'python-docx' library not there."
    
    ensure_output_dir()
    doc = None
    word_doc = Document()
    filename_for_log = "input_stream"
    temp_pdf_path_obj = None
    processed_pages = 0

    # Add basic style (optional)
    style = word_doc.styles['Normal']
    font = style.font
    font.name = 'Calibri' # Or another common font
    font.size = Pt(11)

    try:
        # --- Input Handling (Similar to compress_pdf) ---
        if isinstance(pdf_path_or_stream, (str, Path)):
            pdf_path = str(pdf_path_or_stream)
            filename_for_log = Path(pdf_path).name
            doc = fitz.open(pdf_path)
        elif isinstance(pdf_path_or_stream, (io.BytesIO, io.BufferedReader)):
            filename_for_log = getattr(pdf_path_or_stream, 'filename', 'input_stream')
            temp_dir = OUTPUT_DIR # Save temp in output temporarily
            temp_pdf_path_obj = temp_dir / f"temp_toword_{datetime.now().strftime('%Y%m%d%H%M%S%f')}.pdf"
            logger.info(f"Input stream for PDF-to-Word, saving temp: {temp_pdf_path_obj}")
            with open(temp_pdf_path_obj, 'wb') as f:
                pdf_path_or_stream.seek(0)
                f.write(pdf_path_or_stream.read())
                pdf_path_or_stream.seek(0)
            doc = fitz.open(str(temp_pdf_path_obj))
            filename_for_log = Path(filename_for_log).stem
        else:
            raise TypeError("Unsupported input type for pdf_to_word. Must be path string or stream.")

        if doc.is_encrypted:
            logger.error(f"Cannot convert password-protected PDF to Word: {filename_for_log}")
            if doc: doc.close()
            if temp_pdf_path_obj: cleanup_temp_file(temp_pdf_path_obj)
            return None, f"Error: Input PDF '{filename_for_log}' is password protected."

        logger.info(f"Starting basic PDF-to-Word conversion for '{filename_for_log}'...")

        # --- Process Pages ---
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            logger.debug(f"Processing page {page_num + 1} for text and images...")

            # Extract text blocks with coordinate information
            blocks = page.get_text("dict")["blocks"]
            img_list = page.get_images(full=True)

            # Rough logic: try to place images near where they appear relative to text
            # This is VERY basic and won't handle complex layouts well.
            items = []
            for b in blocks:
                if b['type'] == 0: # Text block
                    for l in b["lines"]:
                        for s in l["spans"]:
                             items.append({'type': 'text', 'bbox': s['bbox'], 'text': s['text']})

            for img_index, img_info in enumerate(img_list):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                img_ext = base_image["ext"]
                # Get image position (optional, might be complex to use accurately)
                # Use page.get_image_bbox(img_info) for coordinates
                img_bbox = page.get_image_bbox(img_info, transform=True) # Get bbox on page
                items.append({'type': 'image', 'bbox': img_bbox, 'bytes': image_bytes, 'ext': img_ext})

            # Sort items approximately by vertical position (top coordinate)
            items.sort(key=lambda item: item['bbox'][1])

            # Add items to Word doc
            for item in items:
                if item['type'] == 'text':
                    word_doc.add_paragraph(item['text'])
                elif item['type'] == 'image':
                    try:
                        img_stream = io.BytesIO(item['bytes'])
                        # Basic width calculation (limit to page width) - adjust as needed
                        img_width_pixels = item['bbox'][2] - item['bbox'][0]
                        page_width_pt = page.rect.width
                        # Estimate Word page width (approx 6 inches usable)
                        word_page_width_inches = 6.0
                        # Scale image width relative to page width
                        scale_factor = img_width_pixels / page_width_pt
                        est_img_width_inches = word_page_width_inches * scale_factor
                        # Prevent excessively large images
                        final_img_width_inches = min(est_img_width_inches, word_page_width_inches)

                        word_doc.add_picture(img_stream, width=Inches(final_img_width_inches))
                        img_stream.close()
                    except Exception as img_err:
                        logger.warning(f"Could not add image from page {page_num + 1} to Word doc: {img_err}")
                        word_doc.add_paragraph(f"[Image Processing Error: {img_err}]")

            # Add a page break (except after the last page)
            if page_num < doc.page_count - 1:
                word_doc.add_page_break()

            processed_pages += 1


        doc.close() # Close the PDF document

        if processed_pages == 0:
             return None, "No pages could be processed from the PDF."

        # --- Save Word Document ---
        output_path = get_output_filename(output_filename_base or Path(filename_for_log).stem, "converted", ".docx")
        word_doc.save(str(output_path))
        logger.info(f"Basic PDF-to-Word conversion saved to: {output_path}")
        return output_path, None

    except ImportError:
         logger.error("python-docx library not found. Please install it (`pip install python-docx`).")
         if doc: doc.close()
         return None, "Error: Required 'python-docx' library not found."
    except Exception as e:
        logger.error(f"Error converting PDF '{filename_for_log}' to Word: {e}", exc_info=True)
        if doc: doc.close()
        return None, f"Error converting PDF to Word: {e}"
    finally:
        if temp_pdf_path_obj:
             cleanup_temp_file(temp_pdf_path_obj)

# --- END PDF TO WORD ---

# Major-Project/pdf_operations.py

import os
import fitz  # PyMuPDF
from fitz.utils import getColor # Import getColor for named colors
from pathlib import Path
import logging
from datetime import datetime

logger = logging.getLogger(__name__)
OUTPUT_DIR = Path("output") # Ensure this is defined as per your project structure

# --- Helper Functions (ensure these exist or adapt) ---
def ensure_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def get_output_filename(base_name, suffix, extension):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S%f")
    safe_base = "".join(c if c.isalnum() or c in ('_', '-') else '_' for c in str(base_name))
    safe_base = safe_base[:100] # Limit base length
    filename = f"{safe_base}_{suffix}_{timestamp}{extension}"
    return OUTPUT_DIR / filename
# --- End Helper Functions ---


# Major-Project/pdf_operations.py

import os
import fitz  # PyMuPDF
from fitz.utils import getColor
from pathlib import Path
import logging
from datetime import datetime

logger = logging.getLogger(__name__)
OUTPUT_DIR = Path("output") # Ensure this is defined as per your project structure

# --- Helper Functions (ensure these exist or adapt) ---
def ensure_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def get_output_filename(base_name, suffix, extension):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S%f")
    safe_base = "".join(c if c.isalnum() or c in ('_', '-') else '_' for c in str(base_name))
    safe_base = safe_base[:100] # Limit base length
    filename = f"{safe_base}_{suffix}_{timestamp}{extension}"
    return OUTPUT_DIR / filename
# --- End Helper Functions ---

# Major-Project/pdf_operations.py

import os
import fitz  # PyMuPDF
from fitz.utils import getColor
from pathlib import Path
import logging
from datetime import datetime

logger = logging.getLogger(__name__)
OUTPUT_DIR = Path("output") # Ensure this is defined

# --- Helper Functions (ensure these exist or adapt) ---
def ensure_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def get_output_filename(base_name, suffix, extension):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S%f")
    safe_base = "".join(c if c.isalnum() or c in ('_', '-') else '_' for c in str(base_name))
    safe_base = safe_base[:100]
    filename = f"{safe_base}_{suffix}_{timestamp}{extension}"
    return OUTPUT_DIR / filename
# --- End Helper Functions ---

# Major-Project/pdf_operations.py

import os
import fitz  # PyMuPDF
from fitz.utils import getColor
from pathlib import Path
import logging
from datetime import datetime

logger = logging.getLogger(__name__)
OUTPUT_DIR = Path("output") # Ensure this is defined as per your project structure

# --- Helper Functions (ensure these exist or adapt) ---
def ensure_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def get_output_filename(base_name, suffix, extension):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S%f")
    safe_base = "".join(c if c.isalnum() or c in ('_', '-') else '_' for c in str(base_name))
    safe_base = safe_base[:100] # Limit base length
    filename = f"{safe_base}_{suffix}_{timestamp}{extension}"
    return OUTPUT_DIR / filename
# --- End Helper Functions ---

def text_to_pdf(text_content: str, output_filename_base="text_document", font_name="cour", font_size=11):
    """
    Generates a PDF from text_content using page.insert_text() line by line.
    Uses black text on a white background. Includes detailed logging for multi-line processing.
    """
    ensure_output_dir()
    # Using a slightly different name to distinguish this version's output if needed.
    output_filename_base_test = f"{output_filename_base}_multiline_debug" # Changed suffix for clarity
    output_path = get_output_filename(output_filename_base_test, "generated", ".pdf")
    doc = None

    black_color = getColor("black")
    white_color = getColor("white")

    logger.info(f"text_to_pdf (multiline_debug): Input text_content (first 200 chars): '{text_content[:200]}...'")
    lines = text_content.split('\n')
    logger.info(f"text_to_pdf (multiline_debug): Split into {len(lines)} lines.")

    if not lines or (len(lines) == 1 and not lines[0].strip()): # Handle empty or effectively empty content
        logger.warning("text_to_pdf (multiline_debug): No actual text lines to process.")
        # Create a blank PDF if no content
        try:
            doc_temp = fitz.open()
            doc_temp.new_page()
            doc_temp.save(str(output_path))
            logger.info(f"PDF (multiline_debug) generated as blank (no content): {output_path}")
            doc_temp.close()
            return output_path, None # Or you could return an error/message
        except Exception as e_blank:
            logger.error(f"text_to_pdf (multiline_debug): Error creating blank PDF: {e_blank}")
            return None, f"Error creating blank PDF: {e_blank}"


    try:
        doc = fitz.open()
        page = doc.new_page()
        page.draw_rect(page.rect, color=white_color, fill=white_color, overlay=False)
        logger.info(f"text_to_pdf (multiline_debug): Created initial page {page.number + 1} of {doc.page_count}")

        margin = 50  # Points
        line_height = font_size * 1.2
        x_coord = margin
        # Start y_cursor for the baseline of the first line of text
        y_cursor = margin + font_size

        for line_num, line in enumerate(lines):
            # Prepare the line for processing
            processed_line = line.strip() # Remove leading/trailing whitespace

            logger.info(f"text_to_pdf (multiline_debug): --- Processing line {line_num + 1}/{len(lines)} ---")
            logger.info(f"text_to_pdf (multiline_debug): Page: {page.number + 1}, Current y_cursor: {y_cursor:.2f}")
            logger.info(f"text_to_pdf (multiline_debug): Line content (stripped): '{processed_line}'")

            # Page break logic: if the current y_cursor is already too far down for this line's baseline
            if y_cursor > (page.rect.height - margin):
                logger.info(f"text_to_pdf (multiline_debug): Page break triggered. y_cursor ({y_cursor:.2f}) > page_height ({page.rect.height}) - margin ({margin}).")
                page = doc.new_page()
                page.draw_rect(page.rect, color=white_color, fill=white_color, overlay=False)
                y_cursor = margin + font_size  # Reset y_cursor for the new page
                logger.info(f"text_to_pdf (multiline_debug): New page {page.number + 1} created. y_cursor reset to {y_cursor:.2f}.")

            text_insertion_point = fitz.Point(x_coord, y_cursor)

            # Only attempt to insert text if the processed line is not empty
            if processed_line:
                insert_len = page.insert_text(
                    text_insertion_point,
                    processed_line,
                    fontname=font_name,
                    fontsize=font_size,
                    color=black_color
                )
                logger.info(f"text_to_pdf (multiline_debug): Inserted '{processed_line[:50]}...' (len: {len(processed_line)}). PyMuPDF wrote {insert_len} chars.")
                if insert_len < len(processed_line):
                    logger.warning(f"text_to_pdf (multiline_debug): Line {line_num + 1}: Not all characters inserted.")
            else:
                logger.info(f"text_to_pdf (multiline_debug): Line {line_num + 1} is empty after strip, skipping text insertion.")
            
            y_cursor += line_height # Move to the next line position for the *next* iteration

        doc.save(str(output_path))
        logger.info(f"PDF (multiline_debug) generated successfully: {output_path} with {doc.page_count} page(s).")
        return output_path, None
    except Exception as e:
        logger.error(f"Error in text_to_pdf (multiline_debug) for '{output_filename_base}': {e}", exc_info=True)
        # Attempt to save even if error occurs partway, for debugging
        if doc and (doc.page_count > 0 or page is not None): # Check if doc or page object exists
            error_file_name = output_path.stem + "_error.pdf"
            error_path = output_path.with_name(error_file_name)
            try:
                doc.save(str(error_path))
                logger.info(f"Saved partial PDF to {error_path} after error.")
            except Exception as save_err:
                logger.error(f"Could not save partial PDF after error: {save_err}")
        return None, f"Error creating PDF (multiline_debug): {e}"
    finally:
        if doc:
            doc.close()

# --- PLACEHOLDERS for PDF to PPT/Excel ---
def pdf_to_powerpoint(pdf_path_or_stream, output_filename_base="converted"):
    """Placeholder function for PDF to PowerPoint."""
    logger.warning("PDF to PowerPoint conversion is highly complex and not implemented.")
    # Basic implementation would involve extracting images/text per page (like pdf_to_word)
    # and adding them to slides using python-pptx. Formatting loss would be significant.
    return None, "Error: PDF to PowerPoint conversion is not supported due to its complexity and likely poor formatting results."

def pdf_to_excel(pdf_path_or_stream, output_filename_base="converted"):
    """Placeholder function for PDF to Excel."""
    logger.warning("PDF to Excel conversion (especially tables) is highly complex and not implemented.")
    # Requires sophisticated table detection (e.g., libraries like camelot-py, tabula-py, or complex image processing/OCR)
    # and then writing data using openpyxl or similar.
    return None, "Error: PDF to Excel conversion (especially tables) is not supported due to its complexity. Consider dedicated table extraction tools."
# --- END PLACEHOLDERS ---

# --- Utility Function (Ensure exists or add if missing from your version) ---
def cleanup_temp_file(filepath):
    """Removes a temporary file if it exists."""
    if filepath and Path(filepath).exists() and Path(filepath).is_file(): # Check it's a file
        try:
            os.remove(filepath)
            logger.info(f"Removed temporary file: {filepath}")
        except OSError as e:
            logger.warning(f"Could not remove temporary file {filepath}: {e}")
    elif filepath:
         logger.debug(f"Cleanup requested but file not found or not a file: {filepath}")